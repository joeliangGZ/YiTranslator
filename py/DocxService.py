import os
import re
import uuid
import uvicorn
import time
from pathlib import Path
from fastapi import FastAPI, UploadFile, File, HTTPException
from docx import Document
from docx.oxml import parse_xml
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from typing import List, Dict

app = FastAPI()

TEMPLATE_DIR = "templates/"
PRODUCT_DIR = "products/"
Path(TEMPLATE_DIR).mkdir(parents=True, exist_ok=True)
Path(PRODUCT_DIR).mkdir(parents=True, exist_ok=True)


class DocumentItem:
    def __init__(self, placeholder_number: int, original_content: str, translate_content: str = None):
        self.placeholder_number = placeholder_number
        self.original_content = original_content
        self.translate_content = translate_content


class DocumentEntity:
    def __init__(self, doc_id: str, original_filename: str, template_filename: str, items: List[DocumentItem]):
        self.doc_id = doc_id
        self.original_filename = original_filename
        self.template_filename = template_filename
        self.items = items


@app.post("/extract/")
async def extract_content(file: UploadFile = File(...)):
    try:
        doc_id = str(uuid.uuid4())
        original_filename = file.filename or "uploaded.docx"
        base_name = original_filename.replace(".docx", "")
        template_filename = f"{base_name}_template_{int(time.time())}.docx"

        # 读取上传文件
        doc = Document(file.file)

        items = []
        current_num = 1

        # 处理段落
        for paragraph in doc.paragraphs:
            current_num = process_paragraph(paragraph, current_num, items)

        # 处理表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        current_num = process_paragraph(paragraph, current_num, items)

        # 保存模板文件
        template_path = os.path.join(TEMPLATE_DIR, template_filename)
        doc.save(template_path)

        return {
            "doc_id": doc_id,
            "original_filename": original_filename,
            "template_filename": template_filename,
            "items": [item.__dict__ for item in items]
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail="文档解析异常")


def process_paragraph(paragraph: Paragraph, current_num: int, items: List[DocumentItem]) -> int:
    original_text = paragraph.text.strip()
    if not original_text:
        return current_num

    # 匹配编号模式
    pattern = re.compile(r"^(\d+(?:\.\d+)*\s+)(.*)$")
    match = pattern.match(original_text)

    if match:
        number_part = match.group(1)
        content = match.group(2)
        new_text = f"{number_part}{{{{ {current_num} }}}}"
    else:
        content = original_text
        new_text = f"{{{{ {current_num} }}}}"

    # 保留原始样式
    if paragraph.runs:
        style = paragraph.runs[0]._element.get_or_add_rPr()
    else:
        style = None

    # 清空段落内容
    paragraph.clear()

    # 添加新内容并保留样式
    run = paragraph.add_run(new_text)
    if style is not None:
        run._element.rPr = style

    items.append(DocumentItem(current_num, content))
    return current_num + 1


@app.post("/fill/")
async def fill_template(doc_entity: dict):
    try:
        template_path = os.path.join(TEMPLATE_DIR, doc_entity["template_filename"])
        if not os.path.exists(template_path):
            raise HTTPException(status_code=404, detail="模板文件未找到")

        content_map = {item["placeholder_number"]: item.get("translate_content")
                       for item in doc_entity["items"]}

        doc = Document(template_path)

        # 处理段落
        for paragraph in doc.paragraphs:
            process_fill_paragraph(paragraph, content_map)

        # 处理表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        process_fill_paragraph(paragraph, content_map)

        # 保存结果文件
        output_filename = doc_entity["original_filename"]
        output_path = os.path.join(PRODUCT_DIR, output_filename)
        doc.save(output_path)

        return {"output_file": output_filename}

    except Exception as e:
        raise HTTPException(status_code=500, detail="文档填充异常")


def process_fill_paragraph(paragraph: Paragraph, content_map: Dict[int, str]):
    original_text = paragraph.text.strip()
    if not original_text:
        return

    # 匹配占位符
    pattern = re.compile(r"\{\{\s*(\d+)\s*\}\}")

    # 保留原始样式
    if paragraph.runs:
        style = paragraph.runs[0]._element.get_or_add_rPr()
    else:
        style = None

    # 替换文本
    new_text = pattern.sub(lambda m: content_map.get(int(m.group(1)), original_text))

    # 清空段落内容
    paragraph.clear()

    # 添加新内容并保留样式
    run = paragraph.add_run(new_text)
    if style is not None:
        run._element.rPr = style


if __name__ == "__main__":
    uvicorn.run("DocxService:app", host="0.0.0.0", port=8000)
